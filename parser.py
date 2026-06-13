"""
parser.py — Resume Text Extraction Module
==========================================
Extracts clean, structured text from PDF and DOCX resume files.

Supported formats
-----------------
- PDF  : via PyMuPDF (fitz) — handles text-native and embedded-font PDFs
- DOCX : via python-docx — preserves paragraph and table content

Usage
-----
    from parser import parse_resume, ParsedResume

    result = parse_resume("resume.pdf")
    print(result.full_text)
    print(result.metadata)

    # Or use format-specific functions directly:
    from parser import extract_pdf_text, extract_docx_text
"""

from __future__ import annotations

import logging
import re
import unicodedata
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


# ── Types ─────────────────────────────────────────────────────────────────────


class FileType(str, Enum):
    PDF = "pdf"
    DOCX = "docx"


@dataclass
class ParsedResume:
    """Structured output returned by every extraction path."""

    file_path: str
    file_type: FileType
    full_text: str                      # Clean, normalised full text
    pages: list[str] = field(default_factory=list)  # Per-page text (PDF only)
    metadata: dict = field(default_factory=dict)     # Doc-level metadata
    warnings: list[str] = field(default_factory=list)

    @property
    def word_count(self) -> int:
        return len(self.full_text.split())

    @property
    def is_empty(self) -> bool:
        return not self.full_text.strip()


class ResumeParseError(Exception):
    """Raised when a file cannot be parsed."""


# ── Text cleaning helpers ──────────────────────────────────────────────────────


def _normalise_unicode(text: str) -> str:
    """Replace ligatures and decompose non-ASCII characters where possible."""
    replacements = {
        "\ufb01": "fi", "\ufb02": "fl",  # fi / fl ligatures
        "\u2019": "'", "\u2018": "'",    # curly apostrophes
        "\u201c": '"', "\u201d": '"',    # curly quotes
        "\u2013": "-", "\u2014": "-",    # en/em dashes
        "\u00a0": " ",                   # non-breaking space
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return unicodedata.normalize("NFKC", text)


def _collapse_whitespace(text: str) -> str:
    """Collapse runs of blank lines to at most two; normalise spaces."""
    # Normalise all line-endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse 3+ consecutive blank lines → 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove trailing whitespace on every line
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return text.strip()


def clean_text(raw: str) -> str:
    """Full cleaning pipeline applied to every extraction result."""
    text = _normalise_unicode(raw)
    text = _collapse_whitespace(text)
    return text


# ── PDF extraction ─────────────────────────────────────────────────────────────


def _extract_page_text(page: fitz.Page) -> str:
    """
    Extract text from a single PyMuPDF page using the 'blocks' strategy,
    which preserves reading order better than the default for most resumes.
    """
    blocks = page.get_text("blocks", sort=True)  # sort=True → reading order
    lines: list[str] = []
    for block in blocks:
        # block = (x0, y0, x1, y1, text, block_no, block_type)
        if block[6] == 0:  # type 0 → text block (type 1 = image)
            text = block[4].strip()
            if text:
                lines.append(text)
    return "\n".join(lines)


def extract_pdf_text(file_path: str | Path) -> ParsedResume:
    """
    Extract text from a PDF resume using PyMuPDF.

    Parameters
    ----------
    file_path : path to the PDF file

    Returns
    -------
    ParsedResume with per-page text and joined full_text

    Raises
    ------
    ResumeParseError on unreadable or encrypted files
    FileNotFoundError if the path does not exist
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    warnings: list[str] = []
    pages_text: list[str] = []
    metadata: dict = {}

    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        raise ResumeParseError(f"Cannot open PDF '{path.name}': {exc}") from exc

    try:
        if doc.is_encrypted:
            # Try with empty password (common for "owner-locked" PDFs)
            if not doc.authenticate(""):
                raise ResumeParseError(
                    f"'{path.name}' is encrypted and requires a password."
                )
            warnings.append("PDF was owner-locked; opened with empty password.")

        # Document metadata
        meta = doc.metadata or {}
        metadata = {
            "title": meta.get("title", ""),
            "author": meta.get("author", ""),
            "creator": meta.get("creator", ""),
            "page_count": doc.page_count,
        }

        for page_num, page in enumerate(doc, start=1):
            raw = _extract_page_text(page)
            if not raw.strip():
                warnings.append(f"Page {page_num} yielded no text (image-only?).")
            pages_text.append(clean_text(raw))

    finally:
        doc.close()

    full_text = clean_text("\n\n".join(p for p in pages_text if p))

    if not full_text.strip():
        warnings.append(
            "No text could be extracted. The PDF may be a scanned image; "
            "consider running an OCR pass first."
        )

    logger.info(
        "PDF parsed: file=%s pages=%d words=%d",
        path.name,
        len(pages_text),
        len(full_text.split()),
    )

    return ParsedResume(
        file_path=str(path),
        file_type=FileType.PDF,
        full_text=full_text,
        pages=pages_text,
        metadata=metadata,
        warnings=warnings,
    )


# ── DOCX extraction ────────────────────────────────────────────────────────────


def _iter_paragraph_text(doc: Document) -> list[str]:
    """
    Yield text from every paragraph in the document body, including
    paragraphs nested inside text boxes (drawing elements).
    """
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also capture text inside drawing / text-box elements (common in resumes)
    for shape in doc.element.body.iter(qn("w:txbxContent")):
        for p_elem in shape.iter(qn("w:p")):
            texts = [
                node.text
                for node in p_elem.iter(qn("w:t"))
                if node.text
            ]
            text = "".join(texts).strip()
            if text:
                paragraphs.append(text)

    return paragraphs


def _iter_table_text(doc: Document) -> list[str]:
    """Extract all non-empty cell text from every table in the document."""
    rows: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            # De-duplicate merged cells (python-docx repeats them)
            seen: list[str] = []
            for c in cells:
                if not seen or seen[-1] != c:
                    seen.append(c)
            if seen:
                rows.append("  |  ".join(seen))
    return rows


def extract_docx_text(file_path: str | Path) -> ParsedResume:
    """
    Extract text from a DOCX resume using python-docx.

    Captures:
    - Body paragraphs (in document order)
    - Text boxes / drawing elements
    - Tables (cell content joined with ' | ')
    - Core document properties as metadata

    Parameters
    ----------
    file_path : path to the .docx file

    Returns
    -------
    ParsedResume (pages list is empty — DOCX has no page concept at parse time)

    Raises
    ------
    ResumeParseError on unreadable files
    FileNotFoundError if the path does not exist
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ResumeParseError(
            f"Cannot open DOCX '{path.name}': {exc}"
        ) from exc

    warnings: list[str] = []

    paragraphs = _iter_paragraph_text(doc)
    tables = _iter_table_text(doc)

    # Interleave: python-docx gives paragraphs and tables separately;
    # we append table rows at the end (ordering them correctly would require
    # parsing the XML body, which is overkill for plain text extraction).
    all_parts = paragraphs + (["--- tables ---"] + tables if tables else [])

    full_text = clean_text("\n".join(all_parts))

    # Core properties
    cp = doc.core_properties
    metadata: dict = {
        "title": cp.title or "",
        "author": cp.author or "",
        "created": str(cp.created) if cp.created else "",
        "modified": str(cp.modified) if cp.modified else "",
    }

    if not full_text.strip():
        warnings.append("No text extracted; the document may be empty or image-only.")

    logger.info(
        "DOCX parsed: file=%s paragraphs=%d table_rows=%d words=%d",
        path.name,
        len(paragraphs),
        len(tables),
        len(full_text.split()),
    )

    return ParsedResume(
        file_path=str(path),
        file_type=FileType.DOCX,
        full_text=full_text,
        pages=[],           # DOCX has no addressable page model
        metadata=metadata,
        warnings=warnings,
    )


# ── Unified entry point ────────────────────────────────────────────────────────


_EXTRACTORS = {
    ".pdf": extract_pdf_text,
    ".docx": extract_docx_text,
}


def parse_resume(file_path: str | Path) -> ParsedResume:
    """
    Detect file type by extension and dispatch to the correct extractor.

    Parameters
    ----------
    file_path : str or Path — absolute or relative path to a PDF or DOCX file

    Returns
    -------
    ParsedResume dataclass

    Raises
    ------
    ValueError          — unsupported file extension
    FileNotFoundError   — file does not exist
    ResumeParseError    — file exists but cannot be parsed
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        supported = ", ".join(_EXTRACTORS)
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {supported}"
        )

    return extractor(path)
