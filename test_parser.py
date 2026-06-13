"""
tests/test_parser.py — Unit tests for parser.py
================================================
Creates minimal in-memory PDF and DOCX fixtures so no external files are needed.
Run with: pytest tests/test_parser.py -v
"""

import textwrap
from io import BytesIO
from pathlib import Path

import fitz
import pytest
from docx import Document

# Adjust import path when running from project root
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from parser import (
    FileType,
    ParsedResume,
    ResumeParseError,
    clean_text,
    extract_docx_text,
    extract_pdf_text,
    parse_resume,
)


# ── Fixtures ───────────────────────────────────────────────────────────────────


@pytest.fixture()
def sample_pdf(tmp_path: Path) -> Path:
    """Create a minimal two-page PDF with known text."""
    pdf_path = tmp_path / "resume.pdf"
    doc = fitz.open()
    for i, content in enumerate(["Alice Smith\nSoftware Engineer", "Skills: Python, FastAPI"], 1):
        page = doc.new_page()
        page.insert_text((72, 72), content, fontsize=12)
    doc.save(str(pdf_path))
    doc.close()
    return pdf_path


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    """Create a minimal DOCX with paragraphs and a table."""
    docx_path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("Bob Jones")
    doc.add_paragraph("Backend Developer")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "5 years"
    table.cell(1, 0).text = "PostgreSQL"
    table.cell(1, 1).text = "3 years"
    doc.save(str(docx_path))
    return docx_path


# ── clean_text ─────────────────────────────────────────────────────────────────


def test_clean_text_collapses_blank_lines():
    raw = "Line one\n\n\n\n\nLine two"
    result = clean_text(raw)
    assert "\n\n\n" not in result
    assert "Line one" in result
    assert "Line two" in result


def test_clean_text_normalises_ligatures():
    assert clean_text("\ufb01le") == "file"
    assert clean_text("pro\ufb01le") == "profile"


def test_clean_text_strips_trailing_spaces():
    result = clean_text("hello   \nworld   ")
    for line in result.split("\n"):
        assert line == line.rstrip()


# ── PDF extraction ─────────────────────────────────────────────────────────────


def test_extract_pdf_returns_parsed_resume(sample_pdf):
    result = extract_pdf_text(sample_pdf)
    assert isinstance(result, ParsedResume)
    assert result.file_type == FileType.PDF


def test_extract_pdf_captures_text(sample_pdf):
    result = extract_pdf_text(sample_pdf)
    assert "Alice Smith" in result.full_text
    assert "Software Engineer" in result.full_text
    assert "Python" in result.full_text


def test_extract_pdf_page_count(sample_pdf):
    result = extract_pdf_text(sample_pdf)
    assert len(result.pages) == 2


def test_extract_pdf_metadata(sample_pdf):
    result = extract_pdf_text(sample_pdf)
    assert "page_count" in result.metadata
    assert result.metadata["page_count"] == 2


def test_extract_pdf_word_count(sample_pdf):
    result = extract_pdf_text(sample_pdf)
    assert result.word_count > 0


def test_extract_pdf_missing_file():
    with pytest.raises(FileNotFoundError):
        extract_pdf_text("/tmp/does_not_exist.pdf")


# ── DOCX extraction ────────────────────────────────────────────────────────────


def test_extract_docx_returns_parsed_resume(sample_docx):
    result = extract_docx_text(sample_docx)
    assert isinstance(result, ParsedResume)
    assert result.file_type == FileType.DOCX


def test_extract_docx_captures_paragraphs(sample_docx):
    result = extract_docx_text(sample_docx)
    assert "Bob Jones" in result.full_text
    assert "Backend Developer" in result.full_text


def test_extract_docx_captures_tables(sample_docx):
    result = extract_docx_text(sample_docx)
    assert "Python" in result.full_text
    assert "PostgreSQL" in result.full_text


def test_extract_docx_no_pages(sample_docx):
    """DOCX parser returns empty pages list — no page model at extraction time."""
    result = extract_docx_text(sample_docx)
    assert result.pages == []


def test_extract_docx_missing_file():
    with pytest.raises(FileNotFoundError):
        extract_docx_text("/tmp/does_not_exist.docx")


# ── Unified parse_resume ───────────────────────────────────────────────────────


def test_parse_resume_routes_pdf(sample_pdf):
    result = parse_resume(sample_pdf)
    assert result.file_type == FileType.PDF


def test_parse_resume_routes_docx(sample_docx):
    result = parse_resume(sample_docx)
    assert result.file_type == FileType.DOCX


def test_parse_resume_unsupported_extension(tmp_path):
    fake = tmp_path / "resume.txt"
    fake.write_text("some text")
    with pytest.raises(ValueError, match="Unsupported file type"):
        parse_resume(fake)


def test_parse_resume_is_empty_false(sample_pdf):
    result = parse_resume(sample_pdf)
    assert not result.is_empty
